import streamlit as st
import speech_recognition as sr
from pydub import AudioSegment
import tempfile
import os
from docx import Document
from io import BytesIO

# Başlık ve Açıklama
st.title("Lycia-Sound-to-Text")
st.write("Bu uygulama, büyük ses dosyalarını küçük parçalara ayırır ve her bir parçayı metne dönüştürerek sonuçları birleştirir.")

# Fonksiyonlar
def convert_to_wav(file_path):
    """MP3 veya MP4A dosyasını WAV formatına dönüştürme."""
    try:
        # Ses dosyasını yükleyin
        audio = AudioSegment.from_file(file_path)
        
        # Geçici bir WAV dosyası için yol oluşturun
        temp_wav = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
        temp_wav_path = temp_wav.name
        
        # Ses ayarlarını yapın (örnekleme oranı ve kanallar)
        audio = audio.set_frame_rate(16000)  # Örnekleme oranını 16kHz olarak ayarla
        audio = audio.set_channels(1)  # Tek kanallı ses (Mono)
        
        # Ses dosyasını WAV formatında kaydet
        audio.export(temp_wav_path, format="wav")
        
        # Geçici dosyayı kapat ve yolunu döndür
        temp_wav.close()
        return temp_wav_path
    
    except Exception as e:
        st.error(f"Ses dönüştürme hatası: {e}")
        return None

# Uygulama mantığı
uploaded_file = st.file_uploader("Ses dosyasını yükleyin (MP3 veya MP4)", type=["mp3", "mp4"])

if uploaded_file is not None:
    # Geçici bir dosyaya ses dosyasını kaydedin
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_file:
        temp_file.write(uploaded_file.read())
        temp_file_path = temp_file.name

    # MP3 veya MP4 dosyasını WAV formatına dönüştür
    wav_file_path = convert_to_wav(temp_file_path)

    if wav_file_path:
        st.success("Ses dosyası başarıyla WAV formatına dönüştürüldü.")
        st.audio(wav_file_path, format="audio/wav")
    else:
        st.error("Ses dosyası dönüştürme başarısız.")


def split_audio(audio, chunk_length):
    """Ses dosyasını belirlenen sürelere göre küçük parçalara bölme."""
    chunks = [audio[i:i + chunk_length * 1000] for i in range(0, len(audio), chunk_length * 1000)]
    return chunks

def recognize_speech_from_chunks(chunks):
    """Her parça için tanıma işlemini başlat ve metne dönüştür."""
    recognizer = sr.Recognizer()
    complete_text = ""
    for i, chunk in enumerate(chunks):
        chunk_file = f"chunk_{i}.wav"
        chunk.export(chunk_file, format="wav")
        st.write(f"{i + 1}. parça işleniyor...")

        # Her bir parça için tanıma işlemini gerçekleştir
        with sr.AudioFile(chunk_file) as source:
            audio_data = recognizer.record(source)

        try:
            text = recognizer.recognize_google(audio_data, language="tr-TR")
            st.write(f"{i + 1}. parça başarıyla dönüştürüldü.")
            complete_text += text + " "
        except sr.UnknownValueError:
            st.warning(f"{i + 1}. parça anlaşılamadı.")
        except sr.RequestError as e:
            st.error(f"Google API Hatası: {e}")

        # Geçici parça dosyasını silme
        os.remove(chunk_file)
    return complete_text

# Kullanıcıdan dosya yüklemesini isteme
uploaded_file = st.file_uploader("Ses veya Video Dosyanızı Yükleyin", type=["wav", "mp3", "mp4"])

# Parça uzunluğunu belirleme
chunk_length = st.slider("Parçalama Süresi (saniye):", min_value=10, max_value=60, value=30, step=10)

# Başlat Butonu
if st.button("Başlat"):
    if uploaded_file is None:
        st.warning("Lütfen bir ses veya video dosyası yükleyin.")
    else:
        st.info("İşlem başlatılıyor, lütfen bekleyin...")

        # Geçici dosya oluşturma ve sesi işleme
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_file:
            temp_file.write(uploaded_file.read())
            temp_file_path = temp_file.name

        # Dosyayı WAV formatına dönüştürme
        st.info("Dosya işleniyor ve WAV formatına dönüştürülüyor...")
        temp_wav_path = convert_to_wav(temp_file_path)

        # WAV dosyasını bölme ve küçük parçalara ayırma
        st.info(f"Ses dosyası {chunk_length} saniyelik parçalara ayrılıyor...")
        audio = AudioSegment.from_file(temp_wav_path)
        chunks = split_audio(audio, chunk_length)

        # Her parça için tanıma işlemini başlatma
        st.info("Parçalar metne dönüştürülüyor...")
        complete_text = recognize_speech_from_chunks(chunks)

        # Tüm parçaları birleştirip gösterme
        st.success("Tüm parçalar başarıyla metne dönüştürüldü.")
        st.text_area("Dönüştürülen Tam Metin:", complete_text, height=400)

        # Word dosyasını oluşturma
        doc = Document()
        doc.add_heading("Dönüştürülmüş Konuşma Metni", 0)
        doc.add_paragraph(complete_text)

        # Word dosyasını indirme bağlantısı oluşturma
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.success("Metinler başarıyla Word dosyasına dönüştürüldü!")
        st.download_button("Metni Word Dosyası Olarak İndir", buffer, "donusturulmus_konusma_metni.docx")

        # Geçici dosyaları temizleme
        os.remove(temp_file_path)
        os.remove(temp_wav_path)
